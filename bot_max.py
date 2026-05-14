import asyncio
import os
from datetime import datetime
from umaxbot import Bot, Dispatcher, types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ========== НАСТРОЙКИ ==========
BOT_TOKEN = "id7731281504_bot"  # Получите в личном кабинете MAX
DATA_FILE = "users_data.txt"
TEMPLATES_FILE = "templates.txt"
EVENT_DATE_FILE = "event_date.txt"
DEADLINE_FILE = "deadline.txt"

OWNER_IDS = [217770759]  # В MAX ID пользователя — число (узнайте через @userinfobot)

ALLOWED_RANKS = [
    "митрополит", "архиепископ", "епископ", "архимандрит",
    "протоиерей", "иерей", "иеромонах", "архидиакон",
    "протодиакон", "диакон", "игумен"
]

CATEGORIES = [
    "Служащее духовенство",
    "Иподиаконы",
    "Служба протокола",
    "Служба обеспечения",
    "Водители",
    "Хор",
    "Ресторанное обслуживание",
    "Гости"
]

# ========== ХРАНИЛИЩЕ ==========
user_data = {}
all_users = []
templates = []
master_event_date = None
deadline = None

# ========== РАБОТА С ФАЙЛАМИ ==========
def load_master_date():
    global master_event_date
    if os.path.exists(EVENT_DATE_FILE):
        with open(EVENT_DATE_FILE, 'r', encoding='utf-8') as f:
            master_event_date = f.read().strip()

def save_master_date(date_str):
    global master_event_date
    master_event_date = date_str
    with open(EVENT_DATE_FILE, 'w', encoding='utf-8') as f:
        f.write(date_str)

def load_deadline():
    global deadline
    if os.path.exists(DEADLINE_FILE):
        with open(DEADLINE_FILE, 'r', encoding='utf-8') as f:
            deadline_str = f.read().strip()
            try:
                deadline = datetime.strptime(deadline_str, "%Y-%m-%d %H:%M")
            except:
                deadline = None

def save_deadline(deadline_dt):
    global deadline
    deadline = deadline_dt
    with open(DEADLINE_FILE, 'w', encoding='utf-8') as f:
        f.write(deadline_dt.strftime("%Y-%m-%d %H:%M"))

def is_collection_active():
    if deadline is None:
        return True
    return datetime.now() < deadline

def load_data():
    global all_users, templates
    all_users = []
    templates = []
    
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    parts = line.strip().split('|')
                    if len(parts) == 6:
                        all_users.append({
                            'surname': parts[0],
                            'name': parts[1],
                            'patronymic': parts[2],
                            'rank': parts[3],
                            'birthdate': parts[4],
                            'category': parts[5]
                        })
    
    if os.path.exists(TEMPLATES_FILE):
        with open(TEMPLATES_FILE, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    parts = line.strip().split('|')
                    if len(parts) == 7:
                        templates.append({
                            'surname': parts[0],
                            'name': parts[1],
                            'patronymic': parts[2],
                            'rank': parts[3],
                            'birthdate': parts[4],
                            'category': parts[5],
                            'used': parts[6] == 'True'
                        })

def save_templates():
    with open(TEMPLATES_FILE, 'w', encoding='utf-8') as f:
        for t in templates:
            f.write(f"{t['surname']}|{t['name']}|{t['patronymic']}|{t['rank']}|{t['birthdate']}|{t['category']}|{t['used']}\n")

def save_data(surname, name, patronymic, rank, birthdate, category):
    with open(DATA_FILE, 'a', encoding='utf-8') as f:
        f.write(f"{surname}|{name}|{patronymic}|{rank}|{birthdate}|{category}\n")
    all_users.append({
        'surname': surname,
        'name': name,
        'patronymic': patronymic,
        'rank': rank,
        'birthdate': birthdate,
        'category': category
    })

def save_template(surname, name, patronymic, rank, birthdate, category):
    templates.append({
        'surname': surname,
        'name': name,
        'patronymic': patronymic,
        'rank': rank,
        'birthdate': birthdate,
        'category': category,
        'used': False
    })
    save_templates()

def mark_template_used(index):
    if 0 <= index < len(templates):
        templates[index]['used'] = True
        save_templates()
        return True
    return False

def reset_templates_usage():
    for t in templates:
        t['used'] = False
    save_templates()

def delete_template(index):
    if 0 <= index < len(templates):
        del templates[index]
        save_templates()
        return True
    return False

def get_sorted_users():
    return sorted(all_users, key=lambda x: x['surname'].lower())

def format_date_for_filename(date_str):
    return date_str.replace('.', '_')

def is_duplicate(surname, name, patronymic, birthdate):
    surname_lower = surname.lower()
    name_lower = name.lower()
    patronymic_lower = patronymic.lower() if patronymic else ""
    
    for user in all_users:
        if (user['surname'].lower() == surname_lower and
            user['name'].lower() == name_lower and
            user['patronymic'].lower() == patronymic_lower and
            user['birthdate'] == birthdate):
            return True
    return False

# ========== СОЗДАНИЕ WORD-ТАБЛИЦ ==========
def create_word_table(event_date):
    doc = Document()
    
    title_text = f'Список участников Богослужения {event_date} (по алфавиту)'
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = '№'
    header_cells[1].text = 'ФИО и должность'
    header_cells[2].text = 'Дата рождения'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    sorted_users = get_sorted_users()
    
    for idx, user in enumerate(sorted_users, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        
        surname_upper = user['surname'].upper()
        name_capital = user['name'].capitalize()
        patronymic_capital = user['patronymic'].capitalize() if user['patronymic'] else ''
        rank = user.get('rank', '')
        
        paragraph = row[1].paragraphs[0]
        
        run_surname = paragraph.add_run(surname_upper + ' ')
        run_surname.font.bold = True
        
        run_name = paragraph.add_run(name_capital)
        run_name.font.bold = False
        if patronymic_capital:
            run_patronymic = paragraph.add_run(' ' + patronymic_capital)
            run_patronymic.font.bold = False
        
        if rank:
            run_rank = paragraph.add_run(f' ({rank})')
            run_rank.font.italic = True
            run_rank.font.bold = False
        
        row[2].text = user['birthdate']
    
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(4)
    table.columns[2].width = Inches(1.5)
    
    doc.add_paragraph()
    stats = doc.add_paragraph(f"Всего записей: {len(all_users)}")
    stats.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    filename = f"users_{format_date_for_filename(event_date)}.docx"
    doc.save(filename)
    return filename

def create_word_table_by_categories(event_date):
    doc = Document()
    
    title_text = f'Список участников Богослужения {event_date} (по категориям)'
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    global_counter = 1
    
    for category in CATEGORIES:
        users_in_category = [
            u for u in all_users 
            if u.get('category', '') == category
        ]
        users_in_category.sort(key=lambda x: x['surname'].lower())
        
        if not users_in_category:
            continue
        
        doc.add_heading(category, level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '№'
        header_cells[1].text = 'ФИО и должность'
        header_cells[2].text = 'Дата рождения'
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for user in users_in_category:
            row = table.add_row().cells
            row[0].text = str(global_counter)
            global_counter += 1
            
            surname_upper = user['surname'].upper()
            name_capital = user['name'].capitalize()
            patronymic_capital = user['patronymic'].capitalize() if user['patronymic'] else ''
            rank = user.get('rank', '')
            
            paragraph = row[1].paragraphs[0]
            
            run_surname = paragraph.add_run(surname_upper + ' ')
            run_surname.font.bold = True
            
            run_name = paragraph.add_run(name_capital)
            run_name.font.bold = False
            if patronymic_capital:
                run_patronymic = paragraph.add_run(' ' + patronymic_capital)
                run_patronymic.font.bold = False
            
            if rank:
                run_rank = paragraph.add_run(f' ({rank})')
                run_rank.font.italic = True
                run_rank.font.bold = False
            
            row[2].text = user['birthdate']
        
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(4)
        table.columns[2].width = Inches(1.5)
        
        doc.add_paragraph()
    
    doc.add_paragraph()
    stats = doc.add_paragraph(f"Всего записей: {len(all_users)}")
    stats.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    filename = f"users_by_categories_{format_date_for_filename(event_date)}.docx"
    doc.save(filename)
    return filename

# ========== КОМАНДЫ ВЛАДЕЛЬЦА ==========
async def cmd_setdate(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может устанавливать дату.")
        return
    
    parts = message.text.split()
    if len(parts) < 2:
        await message.answer(
            "📅 **Использование:**\n"
            "/setdate ДД.ММ.ГГГГ\n\n"
            "Пример: /setdate 25.05.2025"
        )
        return
    
    date_str = parts[1]
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        save_master_date(date_str)
        await message.answer(f"✅ **Дата мероприятия установлена:** {date_str}")
    except ValueError:
        await message.answer("❌ Неверный формат! Используйте ДД.ММ.ГГГГ")

async def cmd_setdeadline(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может устанавливать дедлайн.")
        return
    
    parts = message.text.split()
    if len(parts) < 3:
        await message.answer(
            "⏰ **Использование:**\n"
            "/setdeadline ДД.ММ.ГГГГ ЧЧ:ММ\n\n"
            "Пример: /setdeadline 25.05.2025 18:00"
        )
        return
    
    date_str = parts[1]
    time_str = parts[2]
    
    try:
        deadline_dt = datetime.strptime(f"{date_str} {time_str}", "%d.%m.%Y %H:%M")
        save_deadline(deadline_dt)
        await message.answer(
            f"⏰ **Дедлайн установлен!**\n\n"
            f"📅 Дата и время: {deadline_dt.strftime('%d.%m.%Y %H:%M')}\n\n"
            f"Пользователи смогут добавлять данные только до этого времени."
        )
    except ValueError:
        await message.answer("❌ Неверный формат! Используйте: /setdeadline ДД.ММ.ГГГГ ЧЧ:ММ")

async def cmd_deadline(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может просматривать дедлайн.")
        return
    
    if deadline:
        remaining = deadline - datetime.now()
        if remaining.total_seconds() > 0:
            days = remaining.days
            hours = remaining.seconds // 3600
            minutes = (remaining.seconds % 3600) // 60
            await message.answer(
                f"⏰ **Текущий дедлайн:** {deadline.strftime('%d.%m.%Y %H:%M')}\n\n"
                f"⏳ Осталось: {days} дн. {hours} ч. {minutes} мин."
            )
        else:
            await message.answer(f"⏰ **Дедлайн:** {deadline.strftime('%d.%m.%Y %H:%M')}\n\n❌ Сбор данных завершён!")
    else:
        await message.answer("⚠️ Дедлайн не установлен. Используйте /setdeadline")

async def cmd_add_template(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может создавать шаблоны.")
        return
    
    user_data[user_id] = {"step": "surname", "data": {}, "is_template": True}
    await message.answer("📝 **Создание нового шаблона**\n\nВведите **Фамилию**:")

async def cmd_templates(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может просматривать шаблоны.")
        return
    
    if not templates:
        await message.answer("📭 Нет сохранённых шаблонов.\n\nСоздайте командой /add_template")
        return
    
    msg = "📋 **Список шаблонов:**\n\n"
    for i, t in enumerate(templates, 1):
        status_icon = "🟢" if not t['used'] else "🔵"
        status_text = "активен" if not t['used'] else "использован"
        msg += f"{status_icon} {i}. {t['surname']} {t['name']}"
        if t['patronymic']:
            msg += f" {t['patronymic']}"
        if t['rank']:
            msg += f" ({t['rank']})"
        msg += f"\n   📅 {t['birthdate']} | 📂 {t['category']} | *{status_text}*\n\n"
    
    if len(msg) > 4000:
        msg = msg[:3950] + "\n\n...(список слишком длинный)"
    
    await message.answer(msg)

async def cmd_del_template(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может удалять шаблоны.")
        return
    
    if not templates:
        await message.answer("📭 Нет шаблонов для удаления.")
        return
    
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    for i, t in enumerate(templates):
        name = f"{t['surname']} {t['name']}"
        if t['patronymic']:
            name += f" {t['patronymic']}"
        keyboard.add(types.InlineKeyboardButton(f"{i+1}. {name}", callback_data=f"del_temp_{i}"))
    
    keyboard.add(types.InlineKeyboardButton("❌ Отмена", callback_data="cancel_delete"))
    await message.answer("🗑 Выберите шаблон для удаления:", reply_markup=keyboard)

async def cmd_reset_templates(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Только владелец бота может сбрасывать шаблоны.")
        return
    
    reset_templates_usage()
    await message.answer(f"✅ **Статус всех шаблонов сброшен!**\n\nТеперь все {len(templates)} шаблонов снова активны (🟢).")

async def cmd_stats(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Статистика доступна только владельцу.")
        return
    
    active_templates = len([t for t in templates if not t['used']])
    used_templates = len([t for t in templates if t['used']])
    
    deadline_status = ""
    if deadline:
        if datetime.now() < deadline:
            remaining = deadline - datetime.now()
            days = remaining.days
            hours = remaining.seconds // 3600
            minutes = (remaining.seconds % 3600) // 60
            deadline_status = f"\n⏳ Сбор активен, осталось: {days}д {hours}ч {minutes}м"
        else:
            deadline_status = f"\n✅ Сбор завершён (дедлайн: {deadline.strftime('%d.%m.%Y %H:%M')})"
    else:
        deadline_status = "\n⚠️ Дедлайн не установлен"
    
    await message.answer(
        f"📊 **Статистика бота**\n\n"
        f"Всего записей: {len(all_users)}\n"
        f"Активных сессий: {len(user_data)}\n\n"
        f"**Шаблоны:**\n"
        f"- Всего: {len(templates)}\n"
        f"- 🟢 Активных: {active_templates}\n"
        f"- 🔵 Использованных: {used_templates}\n\n"
        f"📅 Дата мероприятия: {master_event_date or 'не установлена'}{deadline_status}\n\n"
        f"**Категории:**\n" + "\n".join([f"- {cat}: {len([u for u in all_users if u.get('category') == cat])}" for cat in CATEGORIES])
    )

# ========== ОСНОВНЫЕ ОБРАБОТЧИКИ ПОЛЬЗОВАТЕЛЕЙ ==========
async def cmd_start(message: types.Message):
    user_id = message.sender.id
    
    if user_id not in OWNER_IDS and not is_collection_active():
        if deadline:
            await message.answer(
                f"⏰ **Сбор данных завершён!**\n\n"
                f"Дедлайн истёк: {deadline.strftime('%d.%m.%Y %H:%M')}\n\n"
                f"Вы больше не можете добавлять данные."
            )
        else:
            await message.answer("⏰ **Сбор данных завершён!**\n\nВы больше не можете добавлять данные.")
        return
    
    if not master_event_date:
        await message.answer("⚠️ **Дата мероприятия ещё не установлена администратором.**\n\nПожалуйста, подождите.")
        return
    
    user_data[user_id] = {"step": "event_date", "data": {}}
    await message.answer(
        f"📅 **Укажите дату мероприятия**\n\n"
        f"Введите дату в формате **ДД.ММ.ГГГГ**\n"
        f"Пример: 25.05.2025"
    )

async def handle_text(message: types.Message):
    user_id = message.sender.id
    text = message.text.strip()
    
    if user_id not in OWNER_IDS and not is_collection_active():
        if deadline:
            await message.answer(f"⏰ **Сбор данных завершён!** Дедлайн истёк: {deadline.strftime('%d.%m.%Y %H:%M')}")
        else:
            await message.answer("⏰ **Сбор данных завершён!**")
        return
    
    if user_id not in user_data:
        await cmd_start(message)
        return
    
    current_step = user_data[user_id]["step"]
    is_template = user_data[user_id].get("is_template", False)
    
    # Шаг 0: ввод даты мероприятия
    if current_step == "event_date":
        try:
            datetime.strptime(text, "%d.%m.%Y")
            
            if text != master_event_date:
                await message.answer(
                    f"❌ **Ошибка: дата не совпадает!**\n\n"
                    f"Вы ввели: {text}\n\n"
                    f"Пожалуйста, введите правильную дату:"
                )
                return
            
            user_data[user_id]["event_date"] = text
            user_data[user_id]["step"] = "surname"
            user_data[user_id]["data"] = {}
            
            await message.answer(f"✅ **Дата мероприятия подтверждена:** {text}\n\nТеперь введите **Фамилию** участника:")
        except ValueError:
            await message.answer("❌ Неверный формат! Используйте **ДД.ММ.ГГГГ**\nПример: 25.05.2025\n\nПопробуйте ещё раз:")
        return
    
    # Валидация ФИО
    if current_step in ["surname", "name", "patronymic"]:
        if not all(c.isalpha() or c in '- ' for c in text):
            await message.answer("❌ Используйте только буквы, дефис или пробел. Попробуйте ещё раз:")
            return
    
    # Валидация должности
    if current_step == "rank":
        if text.lower() not in ["—", "-", "нет", "пропустить", ""] and text.lower() not in ALLOWED_RANKS:
            ranks_list = ", ".join(ALLOWED_RANKS)
            await message.answer(
                f"❌ Недопустимая должность.\n\n"
                f"Разрешённые должности:\n{ranks_list}\n\n"
                f"Или напишите '—' (если должности нет)"
            )
            return
    
    # Валидация категории
    if current_step == "category":
        try:
            cat_num = int(text)
            if 1 <= cat_num <= len(CATEGORIES):
                text = CATEGORIES[cat_num - 1]
            else:
                await message.answer(f"❌ Введите число от 1 до {len(CATEGORIES)}")
                return
        except ValueError:
            await message.answer(f"❌ Введите номер категории (1-{len(CATEGORIES)})")
            return
    
    # Валидация даты рождения
    if current_step == "birthdate":
        try:
            datetime.strptime(text, "%d.%m.%Y")
        except ValueError:
            await message.answer(
                "❌ Неверный формат! Используйте **ДД.ММ.ГГГГ**\n"
                "Пример: 15.05.1990\n\n"
                "Попробуйте ещё раз:"
            )
            return
    
    # Сохраняем данные
    user_data[user_id]["data"][current_step] = text
    
    steps = ["surname", "name", "patronymic", "rank", "category", "birthdate"]
    current_index = steps.index(current_step)
    
    if current_index < len(steps) - 1:
        next_step = steps[current_index + 1]
        user_data[user_id]["step"] = next_step
        
        prompts = {
            "surname": "Фамилия",
            "name": "Имя",
            "patronymic": "Отчество (если нет, напишите '—')",
            "rank": f"Должность (из списка) или '—':\n{', '.join(ALLOWED_RANKS)}",
            "category": "Выберите категорию (напишите номер):\n" + "\n".join([f"{i+1}. {cat}" for i, cat in enumerate(CATEGORIES)]),
            "birthdate": "Дата рождения в формате ДД.ММ.ГГГГ"
        }
        
        await message.answer(f"📝 {prompts[next_step]}:")
    else:
        data = user_data[user_id]["data"]
        
        patronymic = data.get("patronymic", "")
        if patronymic in ["—", "-", "нет", "пропустить", ""]:
            patronymic = ""
        
        rank = data.get("rank", "")
        if rank in ["—", "-", "нет", "пропустить", ""]:
            rank = ""
        else:
            rank = rank.lower()
        
        category = data.get("category", "")
        
        # Проверка на дубликат
        if not is_template and is_duplicate(data["surname"], data["name"], patronymic, data["birthdate"]):
            dup_msg = (
                f"❌ **Ошибка: такой участник уже есть в таблице!**\n\n"
                f"Данные не добавлены, чтобы избежать дублирования.\n\n"
                f"📋 {data['surname']} {data['name']}"
            )
            if patronymic:
                dup_msg += f" {patronymic}"
            dup_msg += f"\n🎂 {data['birthdate']}\n\n"
            dup_msg += "Пожалуйста, проверьте правильность ввода."
            
            await message.answer(dup_msg)
            del user_data[user_id]
            
            keyboard = types.InlineKeyboardMarkup(row_width=1)
            keyboard.add(
                types.InlineKeyboardButton("➕ Добавить другого участника", callback_data="add_more"),
                types.InlineKeyboardButton("🏁 Закончить", callback_data="finish")
            )
            await message.answer("Что делаем дальше?", reply_markup=keyboard)
            return
        
        if is_template:
            # Проверка дубликата шаблона
            template_exists = False
            for t in templates:
                if (t['surname'].lower() == data["surname"].lower() and
                    t['name'].lower() == data["name"].lower() and
                    t['patronymic'].lower() == patronymic.lower() and
                    t['birthdate'] == data["birthdate"]):
                    template_exists = True
                    break
            
            if template_exists:
                await message.answer(
                    f"❌ **Такой шаблон уже существует!**\n\n"
                    f"📋 {data['surname']} {data['name']} {patronymic}\n"
                    f"🎂 {data['birthdate']}\n\n"
                    f"Шаблон не добавлен."
                )
                del user_data[user_id]
                return
            
            save_template(data["surname"], data["name"], patronymic, rank, data["birthdate"], category)
            await message.answer(
                f"✅ **Шаблон сохранён!**\n\n"
                f"📋 {data['surname']} {data['name']} {patronymic}\n"
                f"🎂 {data['birthdate']}\n"
                f"📂 {category}\n"
                f"📌 {rank if rank else '—'}"
            )
        else:
            save_data(data["surname"], data["name"], patronymic, rank, data["birthdate"], category)
            await message.answer(
                f"✅ **Данные сохранены!**\n\n"
                f"📋 {data['surname']} {data['name']} {patronymic}\n"
                f"🎂 {data['birthdate']}\n"
                f"📂 {category}\n"
                f"📌 {rank if rank else '—'}"
            )
        
        del user_data[user_id]
        
        keyboard = types.InlineKeyboardMarkup(row_width=1)
        keyboard.add(
            types.InlineKeyboardButton("➕ Добавить ещё участника", callback_data="add_more"),
            types.InlineKeyboardButton("🏁 Закончить и выйти", callback_data="finish")
        )
        await message.answer("Что делаем дальше?", reply_markup=keyboard)

# ========== АДМИН-ПАНЕЛЬ И ШАБЛОНЫ ==========
async def cmd_admin(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        await message.answer("❌ Доступ запрещён.")
        return
    
    if not master_event_date:
        await message.answer("⚠️ Сначала установите дату мероприятия: /setdate")
        return
    
    # Сначала показываем шаблоны
    if templates:
        keyboard = types.InlineKeyboardMarkup(row_width=1)
        for i, t in enumerate(templates):
            name = f"{t['surname']} {t['name']}"
            if t['patronymic']:
                name += f" {t['patronymic']}"
            if t['rank']:
                name += f" ({t['rank']})"
            
            if t['used']:
                button_text = f"🔵 {name} (использован)"
                callback = f"already_used_{i}"
            else:
                button_text = f"🟢 {name}"
                callback = f"use_temp_{i}"
            
            keyboard.add(types.InlineKeyboardButton(button_text, callback_data=callback))
        
        keyboard.add(types.InlineKeyboardButton("⏭ Пропустить и перейти к таблицам", callback_data="skip_templates"))
        keyboard.add(types.InlineKeyboardButton("🔄 Сбросить статус всех шаблонов", callback_data="reset_templates"))
        
        await message.answer(
            f"📋 **Доступные шаблоны**\n\n"
            f"📅 Дата мероприятия: {master_event_date}\n\n"
            f"🟢 Зелёные — можно добавить\n"
            f"🔵 Синие — уже добавлены\n\n"
            f"_Шаблоны не удаляются, только отмечаются как использованные._",
            reply_markup=keyboard
        )
    else:
        await show_admin_panel(message)

async def show_admin_panel(message: types.Message):
    user_id = message.sender.id
    if user_id not in OWNER_IDS:
        return
    
    deadline_status = ""
    if deadline:
        if datetime.now() < deadline:
            deadline_status = f"\n⏳ Сбор активен до: {deadline.strftime('%d.%m.%Y %H:%M')}"
        else:
            deadline_status = f"\n✅ Сбор завершён (дедлайн: {deadline.strftime('%d.%m.%Y %H:%M')})"
    else:
        deadline_status = "\n⚠️ Дедлайн не установлен"
    
    keyboard = types.InlineKeyboardMarkup(row_width=2)
    keyboard.add(
        types.InlineKeyboardButton("📊 Общая таблица", callback_data="get_word"),
        types.InlineKeyboardButton("📂 По категориям", callback_data="get_word_by_cat"),
        types.InlineKeyboardButton("📋 Список участников", callback_data="list_all"),
        types.InlineKeyboardButton("📝 Управление шаблонами", callback_data="manage_templates"),
        types.InlineKeyboardButton("🔄 Сбросить статус", callback_data="reset_templates")
    )
    
    await message.answer(
        f"🔐 **Панель администратора**\n\n"
        f"📅 Дата мероприятия: **{master_event_date}**{deadline_status}\n"
        f"📋 Шаблонов: {len(templates)} (активных: {len([t for t in templates if not t['used']])})\n"
        f"👥 Участников: {len(all_users)}",
        reply_markup=keyboard
    )

async def show_manage_templates(message: types.Message):
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    keyboard.add(
        types.InlineKeyboardButton("➕ Создать шаблон", callback_data="create_template"),
        types.InlineKeyboardButton("📋 Список шаблонов", callback_data="list_templates"),
        types.InlineKeyboardButton("🗑 Удалить шаблон", callback_data="delete_template"),
        types.InlineKeyboardButton("🔄 Сбросить статус всех", callback_data="reset_templates"),
        types.InlineKeyboardButton("◀️ Назад", callback_data="back_to_admin")
    )
    await message.answer("📝 **Управление шаблонами**", reply_markup=keyboard)

# ========== ОБРАБОТЧИК КНОПОК (Callback) ==========
async def callback_handler(callback_query: types.CallbackQuery):
    user_id = callback_query.sender.id
    data = callback_query.data
    
    # Проверка прав
    if data in ["get_word", "get_word_by_cat", "list_all", "manage_templates", "reset_templates"] and user_id not in OWNER_IDS:
        await callback_query.answer("❌ У вас нет прав для этого действия.")
        return
    
    # Удаление шаблона
    if data.startswith("del_temp_"):
        index = int(data.split("_")[2])
        if delete_template(index):
            await callback_query.answer("✅ Шаблон удалён!")
            await callback_query.message.edit_text("✅ Шаблон удалён!")
        else:
            await callback_query.answer("❌ Ошибка при удалении.")
        return
    
    # Использование шаблона
    if data.startswith("use_temp_"):
        index = int(data.split("_")[2])
        if 0 <= index < len(templates):
            t = templates[index]
            if not t['used']:
                if is_duplicate(t['surname'], t['name'], t['patronymic'], t['birthdate']):
                    await callback_query.answer(f"❌ {t['surname']} {t['name']} уже есть в таблице!")
                    return
                
                save_data(t['surname'], t['name'], t['patronymic'], t['rank'], t['birthdate'], t['category'])
                mark_template_used(index)
                await callback_query.answer(f"✅ Добавлен: {t['surname']} {t['name']}")
                await callback_query.message.edit_text(f"✅ Добавлен: {t['surname']} {t['name']}\n\nШаблон отмечен как использованный (синий).")
                
                # Показываем обновлённый список
                if templates:
                    keyboard = types.InlineKeyboardMarkup(row_width=1)
                    for i, new_t in enumerate(templates):
                        name = f"{new_t['surname']} {new_t['name']}"
                        if new_t['patronymic']:
                            name += f" {new_t['patronymic']}"
                        if new_t['rank']:
                            name += f" ({new_t['rank']})"
                        
                        if new_t['used']:
                            button_text = f"🔵 {name} (использован)"
                            callback = f"already_used_{i}"
                        else:
                            button_text = f"🟢 {name}"
                            callback = f"use_temp_{i}"
                        
                        keyboard.add(types.InlineKeyboardButton(button_text, callback_data=callback))
                    
                    keyboard.add(types.InlineKeyboardButton("⏭ Пропустить и перейти к таблицам", callback_data="skip_templates"))
                    keyboard.add(types.InlineKeyboardButton("🔄 Сбросить статус всех шаблонов", callback_data="reset_templates"))
                    
                    await callback_query.message.answer(
                        f"📋 Обновлённый список шаблонов:\n\n"
                        f"🟢 Зелёные — можно добавить\n"
                        f"🔵 Синие — уже добавлены\n\n"
                        f"_Осталось активных: {len([t for t in templates if not t['used']])}_",
                        reply_markup=keyboard
                    )
                else:
                    await callback_query.message.answer("📭 Все шаблоны использованы!")
                    await show_admin_panel(callback_query.message)
            else:
                await callback_query.answer(f"⚠️ Шаблон {t['surname']} {t['name']} уже был использован.")
        return
    
    if data == "already_used":
        await callback_query.answer("ℹ️ Этот шаблон уже был добавлен на текущее мероприятие.")
        return
    
    if data == "skip_templates":
        await show_admin_panel(callback_query.message)
        return
    
    if data == "reset_templates":
        reset_templates_usage()
        await callback_query.answer("✅ Статус всех шаблонов сброшен!")
        await callback_query.message.edit_text(
            f"✅ **Статус всех шаблонов сброшен!**\n\n"
            f"Теперь все {len(templates)} шаблонов снова активны (🟢)."
        )
        await show_admin_panel(callback_query.message)
        return
    
    if data == "manage_templates":
        await show_manage_templates(callback_query.message)
        return
    
    if data == "create_template":
        user_data[user_id] = {"step": "surname", "data": {}, "is_template": True}
        await callback_query.message.answer(
            "📝 **Создание нового шаблона**\n\n"
            "Введите **Фамилию**:"
        )
        return
    
    if data == "list_templates":
        if not templates:
            await callback_query.message.answer("📭 Нет сохранённых шаблонов.")
            return
        
        msg = "📋 **Список шаблонов:**\n\n"
        for i, t in enumerate(templates, 1):
            status_icon = "🟢" if not t['used'] else "🔵"
            status_text = "активен" if not t['used'] else "использован"
            msg += f"{status_icon} {i}. {t['surname']} {t['name']}"
            if t['patronymic']:
                msg += f" {t['patronymic']}"
            if t['rank']:
                msg += f" ({t['rank']})"
            msg += f"\n   📅 {t['birthdate']} | 📂 {t['category']} | *{status_text}*\n\n"
        
        if len(msg) > 4000:
            msg = msg[:3950] + "\n\n...(список слишком длинный)"
        
        await callback_query.message.answer(msg)
        return
    
    if data == "delete_template":
        if not templates:
            await callback_query.message.answer("📭 Нет шаблонов для удаления.")
            return
        
        keyboard = types.InlineKeyboardMarkup(row_width=1)
        for i, t in enumerate(templates):
            name = f"{t['surname']} {t['name']}"
            if t['patronymic']:
                name += f" {t['patronymic']}"
            keyboard.add(types.InlineKeyboardButton(f"{i+1}. {name}", callback_data=f"del_temp_{i}"))
        
        keyboard.add(types.InlineKeyboardButton("◀️ Назад", callback_data="manage_templates"))
        await callback_query.message.answer("🗑 Выберите шаблон для удаления:", reply_markup=keyboard)
        return
    
    if data == "back_to_admin":
        await show_admin_panel(callback_query.message)
        return
    
    if data == "add_more":
        if master_event_date:
            user_data[user_id] = {"step": "surname", "data": {}, "event_date": master_event_date}
            await callback_query.message.answer("Введите **Фамилию** следующего участника:")
        else:
            await callback_query.message.answer("Сначала установите дату мероприятия через /setdate")
        return
    
    if data == "finish":
        await callback_query.message.answer(
            "👋 Спасибо за работу с ботом!\n\n"
            "Если нужно добавить участников для новой даты, отправьте /start"
        )
        if user_id in user_data:
            del user_data[user_id]
        return
    
    if data == "get_word":
        if not all_users:
            await callback_query.answer("❌ Нет данных для экспорта.")
            return
        
        filename = create_word_table(master_event_date)
        with open(filename, 'rb') as file:
            await callback_query.message.answer_document(
                document=file,
                filename=filename,
                caption=f"📊 Список участников Богослужения {master_event_date} (по алфавиту)"
            )
        os.remove(filename)
        await callback_query.message.answer("✅ Файл отправлен!")
    
    elif data == "get_word_by_cat":
        if not all_users:
            await callback_query.answer("❌ Нет данных для экспорта.")
            return
        
        filename = create_word_table_by_categories(master_event_date)
        with open(filename, 'rb') as file:
            await callback_query.message.answer_document(
                document=file,
                filename=filename,
                caption=f"📂 Список участников Богослужения {master_event_date} (по категориям)"
            )
        os.remove(filename)
        await callback_query.message.answer("✅ Файл отправлен!")
    
    elif data == "list_all":
        if not all_users:
            await callback_query.answer("❌ Нет данных для отображения.")
            return
        
        msg = f"📋 **Список участников ({master_event_date}):**\n\n"
        sorted_users = get_sorted_users()
        for i, user in enumerate(sorted_users, 1):
            msg += f"{i}. {user['surname']} {user['name']}"
            if user['patronymic']:
                msg += f" {user['patronymic']}"
            if user.get('rank'):
                msg += f" ({user['rank']})"
            msg += f"\n   📅 {user['birthdate']} | 📂 {user.get('category', '—')}\n\n"
        
        if len(msg) > 4000:
            msg = msg[:3950] + "\n\n...(список слишком длинный, используйте экспорт в Word)"
        
        await callback_query.message.answer(msg)

# ========== ЗАПУСК ==========
async def main():
    load_master_date()
    load_deadline()
    load_data()
    
    bot = Bot(token=BOT_TOKEN)
    dp = Dispatcher(bot)
    
    # Регистрация команд
    dp.register_message_handler(cmd_start, commands=['start'])
    dp.register_message_handler(cmd_setdate, commands=['setdate'])
    dp.register_message_handler(cmd_setdeadline, commands=['setdeadline'])
    dp.register_message_handler(cmd_deadline, commands=['deadline'])
    dp.register_message_handler(cmd_add_template, commands=['add_template'])
    dp.register_message_handler(cmd_templates, commands=['templates'])
    dp.register_message_handler(cmd_del_template, commands=['del_template'])
    dp.register_message_handler(cmd_reset_templates, commands=['reset_templates'])
    dp.register_message_handler(cmd_stats, commands=['stats'])
    dp.register_message_handler(cmd_admin, commands=['admin'])
    dp.register_message_handler(handle_text)
    dp.register_callback_query_handler(callback_handler)
    
    # ОБЯЗАТЕЛЬНО: удаляем webhook перед запуском
    await bot.delete_webhook()
    
    print("=" * 50)
    print("🤖 БОТ ДЛЯ MAX ЗАПУЩЕН!")
    print(f"📁 Данных в таблице: {len(all_users)}")
    print(f"📋 Шаблонов: {len(templates)}")
    print(f"👑 Владелец ID: {OWNER_IDS}")
    print(f"📅 Дата мероприятия: {master_event_date or 'не установлена'}")
    if deadline:
        if datetime.now() < deadline:
            print(f"⏰ Сбор активен до: {deadline.strftime('%d.%m.%Y %H:%M')}")
        else:
            print(f"⏰ Сбор завершён (дедлайн: {deadline.strftime('%d.%m.%Y %H:%M')})")
    else:
        print(f"⏰ Дедлайн не установлен")
    print("=" * 50)
    
    await dp.start_polling()

if __name__ == "__main__":
    asyncio.run(main())